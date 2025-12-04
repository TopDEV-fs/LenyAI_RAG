from fastapi import APIRouter, Query, WebSocket, WebSocketDisconnect
from fastapi.responses import FileResponse
from datetime import datetime
from typing import Dict, Any
import re
from docx import Document
from app.services.medgemma import generate_medical_message, generate_prompt
from app.services.ingest_pmc import semantic_search
from app.services.constant import questions, prompt_for_chatgpt,system_prompt_for_compare, prompt_for_compare
from app.core.config import OPENAI_API_KEY
from openai import OpenAI
import os

openai_client = OpenAI(api_key=OPENAI_API_KEY)

router = APIRouter()

router = APIRouter()

doc = Document()
@router.get("/chat")
async def chat_endpoint(
    query: str = Query(..., description="The query to search for"),
    top_k: int = Query(4, description="The number of results to return")
) :
    evidences = semantic_search(query, top_k=top_k)
    prompt = generate_prompt(evidences, query)
    response = generate_medical_message(prompt)
    result = {
        "response": response,
        "evidences": evidences,
        "timestamp": datetime.utcnow().isoformat() + "Z"
    }
    return result

@router.get("/compare")
async def compare_endpoint():
    doc.add_heading('Compare result of Leny vs OpenEvidence', 0)
    for i, question in enumerate(questions):
        print(f"Processing Question {i+1}: {question}")
        evidences = semantic_search(question, top_k=4)
        prompt = generate_prompt(evidences, question)
        response_leny = generate_medical_message(prompt)
        response_openevidence_style = openai_client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": prompt_for_chatgpt},
                {"role": "user", "content": question}
            ],
            temperature=0.2,
            max_tokens=1000,
        ).choices[0].message.content
        completed_prompt_for_compare = prompt_for_compare + f"\n\nQuestion {i+1}: {question}\nAnswer of Leny: {response_leny}\nAnswer of OpenEvidence: {response_openevidence_style}\n Prompt of Leny : {prompt}\n"
        compare_result = openai_client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": system_prompt_for_compare},
                {"role": "user", "content": completed_prompt_for_compare}
            ],
            temperature=0.2,
            max_tokens=1000,
        ).choices[0].message.content
        cleaned_text = re.sub(r'\\n[0-9]*', '', compare_result)  # Remove \n and any numbers following it
        cleaned_text = re.sub(r'\\n', ' ', cleaned_text)  # Replace remaining \n with a space
        cleaned_text = re.sub(r'["{}]', '', cleaned_text)  # Remove curly braces and quotes
        final_text = f"{i+1}. {question}\n Answer of Leny\n{response_leny} \n Answer of OpenEvidence\n {response_openevidence_style}\n{cleaned_text}"

        doc.add_paragraph(final_text)
    doc.save("compare_result.docx")
    return FileResponse("compare_result.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="compare_result.docx")

# @router.websocket("/ws/compare-progress")
# async def compare_progress_ws(websocket: WebSocket):
#     await websocket.accept()
#     doc = Document()
#     doc.add_heading('Compare result of Leny vs OpenEvidence', 0)
#     try:
#         total = len(questions)
#         for i, question in enumerate(questions):
#             evidences = semantic_search(question, top_k=4)
#             prompt = generate_prompt(evidences, question)
#             response_leny = generate_medical_message(prompt)
#             response_openevidence_style = openai_client.chat.completions.create(
#                 model="gpt-4",
#                 messages=[
#                     {"role": "system", "content": prompt_for_chatgpt},
#                     {"role": "user", "content": question}
#                 ],
#                 temperature=0.2,
#                 max_tokens=1000,
#             ).choices[0].message.content
#             completed_prompt_for_compare = prompt_for_compare + f"\n\nQuestion {i+1}: {question}\nAnswer of Leny: {response_leny}\nAnswer of OpenEvidence: {response_openevidence_style}\n Prompt of Leny : {prompt}\n"
#             compare_result = openai_client.chat.completions.create(
#                 model="gpt-4",
#                 messages=[
#                     {"role": "system", "content": system_prompt_for_compare},
#                     {"role": "user", "content": completed_prompt_for_compare}
#                 ],
#                 temperature=0.2,
#                 max_tokens=1000,
#             ).choices[0].message.content
#             cleaned_text = re.sub(r'\\n[0-9]*', '', compare_result)
#             cleaned_text = re.sub(r'\\n', ' ', cleaned_text)
#             cleaned_text = re.sub(r'["{}]', '', cleaned_text)
#             final_text = f"{i+1}. {question}\n Answer of Leny\n{response_leny} \n Answer of OpenEvidence\n {response_openevidence_style}\n{cleaned_text}"
#             doc.add_paragraph(final_text)
#             progress = int((i + 1) / total * 100)
#             await websocket.send_json({
#                 "progress": progress,
#                 "current": i + 1,
#                 "question": question,
#             })
#         doc.save("compare_result.docx")
#         await websocket.send_json({"progress": 100, "status": "done", "download_url": "/compare"})
#         await websocket.close()
#     except WebSocketDisconnect:
#         print("Client disconnected")